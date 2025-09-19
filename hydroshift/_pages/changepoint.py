"""A tool to identify changepoints in hydrologic timeseries."""

from pydoc import doc
import traceback
import uuid
from collections import defaultdict
from dataclasses import dataclass, field
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from plotly.graph_objects import Figure

from hydroshift.consts import (
    CP_F1_CAPTION,
    CP_F2_CAPTION,
    CP_T1_CAPTION,
    CP_T2_CAPTION,
    MAX_CACHE_ENTRIES,
    METRICS,
    VALID_ARL0S,
)
from hydroshift.errors import GageNotFoundException
from hydroshift.utils.changepoint import cp_pvalue_batch, cpm_process_stream
from hydroshift.utils.common import num_2_word
from hydroshift.utils.data_retrieval import Gage
from hydroshift.utils.ffa import LP3Analysis
from hydroshift.utils.jinja import render_template, write_template
from hydroshift.utils.plots import combo_cpm, plot_lp3


@dataclass
class ChangePointAnalysis:
    """OOP representation of the changepoint analysis."""

    gage: Gage = field(default_factory=Gage)
    pval_df: dict = None
    cp_dict: dict = field(default_factory=dict)
    ffa_plot = None
    ffa_df: dict = field(default_factory=pd.DataFrame)

    @property
    def nonstationary(self) -> bool:
        """Whether or not a nonstationarity was identified."""
        # Update later
        return len(self.cp_dict) > 0

    @property
    def evidence_level(self) -> str:
        """Text representation of how many tests identified changepoints."""
        test_count = 0
        for cp in self.cp_dict:
            tmp = len(self.cp_dict[cp].split(","))
            test_count = max(test_count, tmp)
        if test_count == 0:
            return "no"
        elif test_count == 1:
            return "limited"
        elif test_count == 2:
            return "moderate"
        elif test_count > 2:
            return "strong"

    def get_change_windows(self, max_dist: float = 10) -> tuple:
        """Find groups of cps that fall within a certain window of one another."""
        groups = []
        test_counts = []
        dates = list(self.cp_dict.keys())
        test_dict = {k: v.split(",") for k, v in self.cp_dict.items()}
        current_group = [dates[0]]
        current_group_tests = set(test_dict[dates[0]])

        for i in range(1, len(dates)):
            if ((dates[i] - current_group[-1]).days / 365) < max_dist:  # 10 years in days
                current_group.append(dates[i])
                current_group_tests = current_group_tests.union(test_dict[dates[i]])
            else:
                groups.append(current_group)
                test_counts.append(len(current_group_tests))
                current_group = [dates[i]]
                current_group_tests = set(test_dict[dates[i]])

        if current_group:
            groups.append(current_group)
            test_counts.append(len(current_group_tests))

        return groups, test_counts

    def get_max_pvalue(self) -> tuple[float, int]:
        """Get the minimum p value where all tests agree and count how often that occurred."""
        all_tests = self.pval_df.fillna(1).max(axis=1).to_frame(name="pval")
        all_tests["run"] = ((all_tests != all_tests.shift(1)) * 1).cumsum()
        for ind, r in all_tests.groupby("pval").agg({"run": pd.Series.nunique}).sort_index().iterrows():
            if r.run > 0:
                min_p = ind
                count = r.run
                break
        return min_p, count

    @property
    def summary_plot(self) -> Figure:
        """Create plotly plot to summarize analysis."""
        return combo_cpm(self.gage.ams, self.pval_df, self.cp_dict)

    @property
    def summary_png(self) -> BytesIO:
        """Export summary plot to png in memory."""
        bio = BytesIO()
        self.summary_plot.write_image(file=bio, width=1100, height=600)
        return bio

    @property
    def ffa_png(self):
        """Export FFA plot to png in memory."""
        bio = BytesIO()
        self.ffa_plot.write_image(file=bio, width=1100, height=600)
        return bio

    @property
    def cp_df(self):
        """Get a dataframe representing changepoints identified in the streaming analysis."""
        cpa_df = pd.DataFrame.from_dict(self.cp_dict, orient="index", columns=["Tests Identifying Change"])
        cpa_df.index = cpa_df.index.date
        return cpa_df

    @property
    def title(self) -> str:
        """Descriptive title of this analysis."""
        return f"Changepoint Analysis for USGS Gage {st.session_state.gage_id}"

    @property
    def summary_text(self) -> str:
        """A text summary of the changepoint analysis."""
        payload = {
            "evidence_level": self.evidence_level,
            "gage_id": self.gage.gage_id,
            "arl0": st.session_state.arlo_slider,
            "burn_in": st.session_state.burn_in,
            "cp_count": num_2_word(len(self.cp_dict)),
            "plural": len(self.cp_dict) != 1,
            "nonstationary": self.nonstationary,
        }
        return render_template("changepoint_summary_1.md", payload)

    @property
    def test_description(self):
        """Analysis methodology."""
        payload = {
            "alr0": st.session_state.arlo_slider,
            "burn_in": st.session_state.burn_in,
        }
        return render_template("changepoint_description.md", payload)

    @property
    def results_text(self) -> str:
        """A detailed description of the results."""
        # Gather stats
        min_p, p_count = self.get_max_pvalue()
        if min_p == 0.001:
            evidence = "strong"
        elif min_p == 0.005:
            evidence = "moderate"
        elif self.pval_df.isna().all().all():
            evidence = "no"
        elif min_p <= 1:
            evidence = "minor"
        groups, _ = self.get_change_windows()

        # format
        payload = {
            "evidence": evidence,
            "min_p": min_p,
            "p_count": num_2_word(p_count),
            "plural": p_count > 1,
            "len_cp": len(self.cp_dict),
            "len_cp_str": num_2_word(len(self.cp_dict)),
            "test_count": num_2_word(len(self.cp_dict[next(iter(self.cp_dict))].split(","))),
            "grp_count": num_2_word(len(groups)),
            "plural_2": len(groups) > 0,
        }
        return render_template("changepoint_summary_2.md", payload)

    @property
    def ffa_text(self) -> str:
        """Methodology for the modified FFA."""
        return render_template("changepoint_ffa.md")

    @property
    def references(self) -> str:
        """Citations."""
        return render_template("changepoint_references.md")

    @property
    def word_data(self) -> BytesIO:
        """Export text as MS word."""
        document = Document()
        s = document.sections[0]
        s.header_distance = Inches(0.2)
        p = s.header.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r = p.add_run()
        r.add_picture('hydroshift/images/dewberry_full_logo.jpg', height=Inches(0.3))
        document.add_heading(self.title, level=1)
        document.add_heading("Summary", level=2)
        self.add_markdown_to_doc(document, self.summary_text)
        document.add_picture(self.summary_png, width=Inches(6.5))
        document.add_heading("Changepoint detection method", level=2)
        self.add_markdown_to_doc(document, self.test_description)
        if len(self.cp_dict) > 0:
            document.add_heading("Changepoint detection results", level=2)
            self.add_markdown_to_doc(document, self.results_text)
            if len(self.cp_dict) > 1:
                self.add_table_from_df(self.cp_df, document, index_name="Date")

        if self.ffa_df is not None and self.ffa_plot is not None:
            document.add_heading("Modified flood frequency analysis", level=2)
            self.add_markdown_to_doc(document, self.ffa_text)
            document.add_picture(self.ffa_png, width=Inches(6.5))
            self.add_markdown_to_doc(document, "**Figure 2.** Modified flood frequency analysis.")
            self.add_markdown_to_doc(document, "**Table 2.** Modified flood quantiles.")
            self.add_table_from_df(self.ffa_df, document, index_name="Regime Period")

        document.add_heading("References", level=2)
        self.add_markdown_to_doc(document, self.references)

        out = BytesIO()
        document.save(out)
        return out

    def add_table_from_df(self, df: pd.DataFrame, document: Document, index_name: str = None):
        if index_name is not None:
            df = df.copy()
            cols = df.columns
            df[index_name] = df.index
            df = df[[index_name, *cols]]

        t = document.add_table(df.shape[0] + 1, df.shape[1])
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])
        t.style = "Light Grid"

    def add_markdown_to_doc(self, document: Document, text: str):
        """Convert some elements of markdown to word format."""
        text = text.replace("\n        ", "")
        p = document.add_paragraph("")
        bold = text.startswith("**")
        if bold:
            text = text[2:]
        for r in text.split("**"):
            p.add_run(r).bold = bold
            bold = not bold

    def validate_data(self):
        """Cache results of get_ams."""
        data = self.gage.ams

        # Validate
        if data is None:
            return False, "Unable to retrieve data."
        elif len(data) < st.session_state.burn_in:
            st.session_state.valid_data = False
            return False, "Not enough peaks available for analysis. {} peaks found, but burn-in length was {}".format(
                len(data["peak_va"]), st.session_state.burn_in
            )
        else:
            return True, None


def define_variables():
    """Set up page state and get default variables."""
    # Ensure gage valid
    if st.session_state["gage"] is None:
        return False, "USGS gage not found"
    if not Gage(st.session_state["gage_id"]).ams_valid:
        return False, "USGS gage has invalid AMS data."
    # Instantiate analysis class and get data
    if "changepoint" not in st.session_state:
        st.session_state.changepoint = ChangePointAnalysis(st.session_state.gage)
    elif st.session_state.changepoint.gage.gage_id != st.session_state.gage_id:
        st.session_state.changepoint = ChangePointAnalysis(st.session_state.gage)
    return st.session_state.changepoint.validate_data()


def refresh_data_editor():
    """Set a new uuid for the data entry widget."""
    st.session_state.data_editor_key = str(uuid.uuid4())


def make_sidebar():
    """User control for analysis."""
    with st.sidebar:
        st.title("Settings")
        try:
            st.session_state["gage_id"] = st.text_input(
                "Enter USGS Gage Number:",
                st.session_state["gage_id"],
                on_change=refresh_data_editor,
            )
            st.session_state.gage = Gage(st.session_state["gage_id"])
        except GageNotFoundException:
            st.session_state.gage = None
            return
        with st.form("changepoint_params"):
            st.select_slider(
                "False Positive Rate (1 in #)",
                options=VALID_ARL0S,
                value=1000,
                key="arlo_slider",
                label_visibility="visible",
            )

            st.number_input(
                "Burn-in Period",
                0,
                100,
                20,
                key="burn_in",
                label_visibility="visible",
            )
            st.text("Flood Frequency Analysis")
            init_data = pd.DataFrame(columns=["Regime Start", "Regime End"])
            init_data["Regime Start"] = pd.to_datetime(init_data["Regime Start"])
            init_data["Regime End"] = pd.to_datetime(init_data["Regime End"])

            # Make data editor.  Unique key allows for refreshing
            if "data_editor_key" not in st.session_state:
                refresh_data_editor()
            start_config = st.column_config.DateColumn("Regime Start", format="D/M/YYYY")
            end_config = st.column_config.DateColumn("Regime End", format="D/M/YYYY")
            st.data_editor(
                init_data,
                num_rows="dynamic",
                key=st.session_state.data_editor_key,
                column_config={"Regime Start": start_config, "Regime End": end_config},
            )
            st.form_submit_button(label="Run Analysis")

        st.divider()
        write_template("data_sources_side_bar.html")


def run_analysis():
    """Run the change point model analysis."""
    cpa = st.session_state.changepoint
    cpa.pval_df = get_pvalues(cpa.gage.ams)
    cpa.cp_dict = get_changepoints(cpa.gage.ams, st.session_state.arlo_slider, st.session_state.burn_in)


@st.cache_data(max_entries=MAX_CACHE_ENTRIES)
def get_pvalues(data: pd.DataFrame) -> pd.DataFrame:
    """Get pvalue df associated with changepoint analysis."""
    ts = data["peak_va"].values
    pval_df = data[[]].copy()
    for metric in METRICS:
        pval_df[metric] = cp_pvalue_batch(metric, ts)
    return pval_df


@st.cache_data(max_entries=MAX_CACHE_ENTRIES)
def get_changepoints(data: pd.DataFrame, arl0: int, burn_in: int) -> dict:
    """Run the process stream analysis and return changepoints identified."""
    ts = data["peak_va"].values
    cp_dict = defaultdict(list)
    for metric in METRICS:
        stream_res = cpm_process_stream(ts, metric, arl0, burn_in)
        for cp in stream_res["changePoints"]:
            cp_dict[data.index[cp]].append(metric)
    for cp in cp_dict:
        cp_dict[cp] = ", ".join(cp_dict[cp])
    return cp_dict


@st.cache_data(max_entries=MAX_CACHE_ENTRIES)
def ffa_analysis(data: pd.DataFrame, regimes: list):
    """Run multiple flood frequency analyses for different regimes."""
    ffas = []
    for r in regimes:
        if "Regime Start" in r and "Regime End" in r:
            sub = data.loc[r["Regime Start"] : r["Regime End"]].copy()
            peaks = sub["peak_va"].values
            label = f"{r['Regime Start']} - {r['Regime End']}"
            lp3 = LP3Analysis(
                st.session_state.gage_id,
                peaks,
                use_map_skew=False,
                est_method="MOM",
                label=label,
            )
            ffas.append(lp3)
    if len(ffas) > 0:
        ffa_plot = plot_lp3(ffas)
        ffa_df = ffas[0].quantile_df[["Recurrence Interval (years)"]]
        ffa_df = ffa_df.set_index("Recurrence Interval (years)")
        for i in ffas:
            ffa_df[i.label] = i.quantile_df["Discharge (cfs)"].values
        return ffa_plot, ffa_df
    else:
        return None, None


def make_body():
    """Assemble main app body."""
    with st.container(width=850):
        cpa: ChangePointAnalysis = st.session_state.changepoint
        st.title(cpa.title)
        warnings()

        st.header("Summary")
        st.markdown(cpa.summary_text)
        st.plotly_chart(cpa.summary_plot, use_container_width=True)
        st.markdown(CP_F1_CAPTION)
        st.header("Changepoint detection method")
        st.markdown(cpa.test_description)

        if len(cpa.cp_dict) > 0:
            st.header("Changepoint detection results")
            st.markdown(cpa.results_text)
            if len(cpa.cp_dict) > 1:
                st.markdown(CP_T1_CAPTION)
                st.table(cpa.cp_df)

        st.header("Modified flood frequency analysis")
        if len(st.session_state[st.session_state.data_editor_key]["added_rows"]) > 0:
            ffa_plot, ffa_df = ffa_analysis(
                cpa.gage.ams,
                st.session_state[st.session_state.data_editor_key]["added_rows"],
            )
            if ffa_plot is not None and ffa_df is not None:
                st.markdown(cpa.ffa_text)
                cpa.ffa_plot = ffa_plot
                cpa.ffa_df = ffa_df
                st.plotly_chart(ffa_plot, use_container_width=True)
                st.markdown(CP_F2_CAPTION)
                st.markdown(CP_T2_CAPTION)
                st.table(ffa_df)
        else:
            st.info(
                "To run pre- and post-changepoint flood frequency analyses, you can input timeseries ranges (regimes) in the flood frequency table on the sidebar.  You may add as many regimes as you think are appropriate, and the periods may overlap."
            )

        st.header("References")
        st.markdown(cpa.references)

        st.download_button(
            "Download analysis",
            cpa.word_data,
            f"changepoint_analysis_{st.session_state.gage_id}.docx",
        )


def warnings():
    """Print warnings on data validity etc."""
    if st.session_state.changepoint.gage.ams is not None and "peak_va" in st.session_state.changepoint.gage.ams.columns:
        if st.session_state.changepoint.gage.missing_dates_ams:
            st.warning(
                "Missing {} dates between {} and {}".format(
                    len(st.session_state.changepoint.gage.missing_dates_ams),
                    st.session_state.changepoint.gage.ams.index.min(),
                    st.session_state.changepoint.gage.ams.index.max(),
                )
            )


def changepoint():
    """Outline the page."""
    st.set_page_config(page_title="USGS Gage Changepoint Analysis", layout="wide")
    make_sidebar()
    valid, msg = define_variables()
    if valid:
        try:
            run_analysis()
        except Exception as e:
            st.error(f"Changepoint model failed with error: {e}")
            st.error(traceback.format_exc())
        else:
            make_body()
    else:
        st.error(msg)
